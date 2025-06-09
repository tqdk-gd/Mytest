from PyQt5.QtCore import QObject, pyqtSignal, QThread, QUrl
import re
from PyQt5 import QtWidgets, QtCore
import os
import datetime
import glob
import yaml
from concurrent.futures import ThreadPoolExecutor
from PyQt5.QtWidgets import QDialog, QLabel, QLineEdit, QComboBox, QPushButton, QVBoxLayout, QFileDialog
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PyQt5.QtWebEngineWidgets import QWebEngineView

class ScanThread(QThread):
    scan_finished = pyqtSignal(list, dict, str)  # 扫描完成信号，传递漏洞信息、漏洞行信息和扫描目录

    def __init__(self, directory, rules,file_extensions):
        super().__init__()
        self.directory = directory
        self.rules = rules
        self.file_extensions = file_extensions

    def run(self):
        vulnerabilities, lines_with_vulnerabilities = self.scan_directory(self.directory, self.rules)
        self.scan_finished.emit(vulnerabilities, lines_with_vulnerabilities, self.directory)

    def scan_directory(self, directory, rules):
        vulnerabilities = []
        max_workers = min(2 * os.cpu_count(), 80)
        lines_with_vulnerabilities = {}
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = []
            for root, _, files in os.walk(directory):
                for file_name in files:
                    # 根据传入的后缀信息筛选文件
                    if any(file_name.endswith(ext) for ext in self.file_extensions):
                        file_path = os.path.join(root, file_name)
                        future = executor.submit(self.scan_file, file_path, rules)
                        futures.append((file_path, future))
            for file_path, future in futures:
                vulns, vuln_lines = future.result()
                vulnerabilities += vulns
                lines_with_vulnerabilities[file_path] = vuln_lines
        return vulnerabilities, lines_with_vulnerabilities

    def scan_file(self, file_path, rules):
        vulnerabilities = []
        lines_with_vulnerabilities = set()
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            php_code = file.read()
            vulns, vuln_lines = self.scan_code(php_code, rules, file_path)
            vulnerabilities += vulns
            lines_with_vulnerabilities.update(vuln_lines)
        return vulnerabilities, lines_with_vulnerabilities

    def scan_code(self, php_code, rules, file_path):
        vulnerabilities = []
        lines_with_vulnerabilities = set()
        for rule in rules:
            if 'only-regex' in rule:
                vulns, vuln_lines = self.only_regex_match(php_code, rule, file_path)
                vulnerabilities += vulns
                lines_with_vulnerabilities.update(vuln_lines)
            elif 'function-param-regex' in rule:
                vulns, vuln_lines = self.function_param_regex_match(php_code, rule, file_path)
                vulnerabilities += vulns
                lines_with_vulnerabilities.update(vuln_lines)
            elif 'customize-match' in rule:
                vulns, vuln_lines = self.customize_match(php_code, rule, file_path)
                vulnerabilities += vulns
                lines_with_vulnerabilities.update(vuln_lines)
        return vulnerabilities, lines_with_vulnerabilities

    def only_regex_match(self, php_code, rule, file_path):
        vulnerabilities = []
        vuln_lines = set()
        pattern = re.compile(rule['only-regex']['pattern'], re.IGNORECASE)
        exclude_patterns = [re.compile(p, re.IGNORECASE) for p in rule.get('exclude-patterns', [])]
        lines = php_code.split('\n')
        for i, line in enumerate(lines, start=1):
            matches = pattern.finditer(line)
            for match in matches:
                if any(ep.match(match.group()) for ep in exclude_patterns):
                    continue
                code_snippet = line[max(match.start() - 50, 0):min(match.end() + 50, len(line))]
                vulnerabilities.append({
                    'name': rule['name'],
                    'description': rule['description'],
                    'file_path': file_path,
                    'code_snippet': code_snippet
                })
                vuln_lines.add(i)
        return vulnerabilities, vuln_lines

    def function_param_regex_match(self, php_code, rule, file_path):
        vulnerabilities = []
        vuln_lines = set()
        func_pattern = re.compile(rule['function-param-regex']['function_pattern'], re.IGNORECASE)
        param_pattern = re.compile(rule['function-param-regex']['param_pattern'], re.IGNORECASE)
        exclude_functions = [re.compile(p, re.IGNORECASE) for p in rule.get('exclude-functions', [])]
        lines = php_code.split('\n')
        for i, line in enumerate(lines, start=1):
            matches = func_pattern.finditer(line)
            for match in matches:
                if any(ef.match(match.group()) for ef in exclude_functions):
                    continue
                params = param_pattern.findall(match.group())
                for param in params:
                    if self.is_tainted(param):
                        code_snippet = line[max(match.start() - 50, 0):min(match.end() + 50, len(line))]
                        vulnerabilities.append({
                            'name': rule['name'],
                            'description': rule['description'],
                            'file_path': file_path,
                            'code_snippet': code_snippet
                        })
                        vuln_lines.add(i)
        return vulnerabilities, vuln_lines

    def is_tainted(self, param):
        return True

    def customize_match(self, php_code, rule, file_path):
        vulnerabilities = []
        vuln_lines = set()
        custom_pattern = re.compile(rule['customize-match']['pattern'], re.IGNORECASE)
        exclude_patterns = [re.compile(p, re.IGNORECASE) for p in rule.get('exclude-patterns', [])]
        lines = php_code.split('\n')
        for i, line in enumerate(lines, start=1):
            matches = custom_pattern.finditer(line)
            for match in matches:
                if any(ep.match(match.group()) for ep in exclude_patterns):
                    continue
                main_func = globals()[rule['customize-match']['main']]
                params = main_func(match.group())
                if any(self.is_tainted(param) for param in params):
                    code_snippet = line[max(match.start() - 50, 0):min(match.end() + 50, len(line))]
                    vulnerabilities.append({
                        'name': rule['name'],
                        'description': rule['description'],
                        'file_path': file_path,
                        'code_snippet': code_snippet
                    })
                    vuln_lines.add(i)
        return vulnerabilities, vuln_lines

class PushButtonRunScanEvent(QObject):
    text_browser_updated = pyqtSignal(str)  # 自定义信号，用于更新文本浏览器内容

    def __init__(self, listWidget_2, rules,comboBox_language):
        super().__init__()
        self.listWidget_2 = listWidget_2
        self.rules = rules
        self.comboBox_language = comboBox_language
        self.scan_thread = None


    def run_scan(self):
        directory = QtWidgets.QFileDialog.getExistingDirectory(None, "选择扫描目录")
        if directory:
            selected_language = self.comboBox_language.currentText()
            # 根据选择的语言确定文件后缀
            if selected_language == "php":
                file_extensions = ['.php']
            elif selected_language == "java":
                file_extensions = ['.java']
            elif selected_language == "python":
                file_extensions = ['.py']
            elif selected_language == "go":
                file_extensions = ['.go']
            else:
                file_extensions = []

            self.listWidget_2.addItem(f"扫描目录: {directory}")
            self.scan_thread = ScanThread(directory, self.rules, file_extensions)
            self.scan_thread.scan_finished.connect(self.handle_scan_finished)
            self.scan_thread.start()
        else:
            self.listWidget_2.addItem("未选择任何目录")

    def handle_scan_finished(self, vulnerabilities, lines_with_vulnerabilities, directory):
        if vulnerabilities:
            # 弹出对话框让用户选择报告类型
            dialog = QDialog()
            dialog.setWindowTitle("选择报告类型")
            layout = QVBoxLayout()

            label = QLabel("请选择要生成的报告类型:")
            layout.addWidget(label)

            combo_box = QComboBox()
            combo_box.addItems(["HTML 报告", "Word 报告"])
            layout.addWidget(combo_box)

            ok_button = QPushButton("确定")
            ok_button.clicked.connect(dialog.accept)
            layout.addWidget(ok_button)

            dialog.setLayout(layout)
            if dialog.exec_() == QDialog.Accepted:
                report_type = combo_box.currentText()
                timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
                if report_type == "HTML 报告":
                    self.export_to_html(vulnerabilities, lines_with_vulnerabilities, f'./baogao/{timestamp}.html', directory)
                    self.listWidget_2.addItem(f"扫描完成，漏洞报告已导出到 ./baogao/{timestamp}.html")
                    # 使用浏览器打开 HTML 报告
                    self.open_html_report(f'./baogao/{timestamp}.html')
                    # os.system(f'start ./baogao/{timestamp}.html')
                elif report_type == "Word 报告":
                    self.export_to_word(vulnerabilities, lines_with_vulnerabilities, f'./baogao/{timestamp}.docx', directory)
                    self.listWidget_2.addItem(f"扫描完成，漏洞报告已导出到 ./baogao/{timestamp}.docx")
                self.listWidget_2.addItem(f"发现漏洞数量: {len(vulnerabilities)}")
        else:
            self.listWidget_2.addItem("扫描完成，未发现任何漏洞")

    def export_to_html(self, vulnerabilities, lines_with_vulnerabilities, output_file, directory):
        # 定义 CSS 样式
        css_style = """
        <style>
            body {
                font-family: Arial, sans-serif;
                margin: 20px;
                background-color: #f4f4f9;
            }
            h1 {
                color: #333;
                text-align: center;
                margin-bottom: 30px;
            }
            h2 {
                color: #1976d2;
                border-bottom: 2px solid #1976d2;
                padding-bottom: 5px;
                margin-top: 40px;
            }
            .vulnerability {
                background-color: #fff;
                padding: 20px;
                margin-bottom: 20px;
                border-radius: 5px;
                box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
                transition: transform 0.2s ease;
            }
            .vulnerability:hover {
                transform: translateY(-3px);
            }
            p {
                margin: 10px 0;
            }
            strong {
                color: #333;
            }
            hr {
                border: 0;
                border-top: 1px solid #ddd;
                margin: 20px 0;
            }
        </style>
        """
        with open(output_file, 'w', encoding='utf-8') as html_file:
            html_file.write("<html>\n<head>\n<title>代码审计报告</title>\n")
            html_file.write(css_style)
            html_file.write("</head>\n<body>\n")
            if vulnerabilities:
                categorized_vulns = self.categorize_vulnerabilities(vulnerabilities)
                html_file.write("<h1>代码审计报告</h1>\n")
                for category, vulns in categorized_vulns.items():
                    html_file.write(f"<h2>审计目录: {directory}</h2>\n")
                    for i, vulnerability in enumerate(vulns):
                        html_file.write("<div class='vulnerability'>\n")
                        html_file.write(f"<p><strong>漏洞名称:</strong> {vulnerability['name']}</p>\n")
                        html_file.write(f"<p><strong>漏洞描述:</strong> {vulnerability['description']}</p>\n")
                        html_file.write(f"<p><strong>文件路径:</strong> {vulnerability['file_path']}</p>\n")
                        html_file.write(f"<p><strong>漏洞行数:</strong> {', '.join(map(str, lines_with_vulnerabilities.get(vulnerability['file_path'], [])))}</p>\n")
                        html_file.write(f"<p><strong>漏洞代码:</strong> <code>{vulnerability['code_snippet']}</code></p>\n")
                        html_file.write("</div>\n")
                        if i < len(vulns) - 1:
                            html_file.write("<hr>\n")
            else:
                html_file.write("<p style='text-align: center; color: #666;'>未发现任何漏洞。</p>\n")
            html_file.write("</body>\n</html>")

    def export_to_word(self, vulnerabilities, lines_with_vulnerabilities, output_file, directory):
        doc = Document()

        # 添加标题
        title = doc.add_heading('代码审计报告', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if vulnerabilities:
            categorized_vulns = self.categorize_vulnerabilities(vulnerabilities)
            for category, vulns in categorized_vulns.items():
                # 添加类别标题
                doc.add_heading(f'审计目录: {directory}', level=1)
                for vulnerability in vulns:
                    # 添加漏洞信息
                    doc.add_heading('漏洞信息', level=2)
                    doc.add_paragraph(f'漏洞名称: {vulnerability["name"]}')
                    doc.add_paragraph(f'漏洞描述: {vulnerability["description"]}')
                    doc.add_paragraph(f'文件路径: {vulnerability["file_path"]}')
                    doc.add_paragraph(f'漏洞行数: {", ".join(map(str, lines_with_vulnerabilities.get(vulnerability["file_path"], [])))}')
                    doc.add_paragraph(f'漏洞代码: {vulnerability["code_snippet"]}')
                    doc.add_paragraph('-' * 80)
        else:
            doc.add_paragraph('未发现任何漏洞。').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 保存文档
        if not os.path.exists('./baogao'):
            os.makedirs('./baogao')
        doc.save(output_file)

    def categorize_vulnerabilities(self, vulnerabilities):
        categories = {}
        for vuln in vulnerabilities:
            category = vuln.get('category', 'Uncategorized')
            if category not in categories:
                categories[category] = []
            categories[category].append(vuln)
        return categories

    def open_html_report(self, html_path):
        # 创建一个新的窗口，并将其保存为类属性
        self.html_report_window = QtWidgets.QMainWindow()
        self.html_report_window.setWindowTitle("代码审计报告")
        self.html_report_window.resize(1200, 800)

        # 创建 QWebEngineView 实例
        web_view = QWebEngineView()
        web_view.load(QUrl.fromLocalFile(os.path.abspath(html_path)))

        # 将 QWebEngineView 设置为窗口的中心部件
        self.html_report_window.setCentralWidget(web_view)

        # 显示窗口
        self.html_report_window.show()